"""
Word 文档解析模块 - python-docx 文本提取

核心能力：
  1. 提取 .docx / .doc 文件中的所有文本内容（按段落保留结构）
  2. 提取表格内容
  3. 保留标题层级信息

依赖：
  - python-docx (.docx)
  - antiword / textract (.doc) — 可选，.doc 无此库时降级

使用方式：
    from src.app.testcase_agent.word_analyzer import analyze_docx

    result = analyze_docx("/path/to/document.docx", user_text="总结这份文档")
    print(result)  # 完整的文档文本内容
"""

import os
from pathlib import Path
from typing import Optional

# ============================================================
# 配置常量
# ============================================================

MAX_WORD_SIZE = 50 * 1024 * 1024  # 50MB


# ============================================================
# 公共接口
# ============================================================

def analyze_docx(docx_path: str, user_text: str = "") -> Optional[str]:
    """
    解析 .docx 文件，提取完整文本和表格内容。

    Args:
        docx_path: 本地 .docx 文件绝对路径
        user_text: 用户针对此文档的问题（会追加到输出末尾）

    Returns:
        提取的文档文本内容；失败返回 None
    """
    if not os.path.isfile(docx_path):
        print(f"[word_analyzer] 文件不存在: {docx_path}")
        return None

    file_size = os.path.getsize(docx_path)
    if file_size > MAX_WORD_SIZE:
        print(f"[word_analyzer] 文件过大 ({file_size}B > {MAX_WORD_SIZE}B)，跳过")
        return None

    try:
        from docx import Document
        doc = Document(docx_path)
    except ImportError:
        print("[word_analyzer] python-docx 未安装 (pip install python-docx)，尝试降级方案")
        return _try_antiword(docx_path)

    sections = []

    # 提取段落文本（按顺序）
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            style_name = para.style.name if para.style else ""
            # 检测是否为标题样式
            is_heading = any(
                style_name.startswith(f"Heading {i}") or style_name.startswith(f"标题 {i}")
                for i in range(1, 10)
            )
            if is_heading:
                sections.append(f"\n## {text}\n")
            else:
                sections.append(text)

    # 提取表格内容
    for table_idx, table in enumerate(doc.tables, start=1):
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(" | ".join(cells))
        if rows_data:
            header_line = rows_data[0] if len(rows_data) > 0 else ""
            separator = " | ".join(["---"] * len(header_line.split("|"))).strip()
            sections.append(
                f"\n### 表格 {table_idx}\n"
                f"| {header_line} |\n"
                f"| {separator} |\n"
                + "\n".join([f"| {row} |" for row in rows_data])
            )

    if not sections:
        print(f"[word_analyzer] ⚠️ .docx 内容为空: {docx_path}")
        return None

    full_text = "\n\n".join(sections)

    if user_text:
        full_text += f"\n\n> 用户问题: {user_text}"

    print(f"[word_analyzer] ✅ .docx 解析完成 | 路径={docx_path}"
                 f" | 段落数={len(doc.paragraphs)} | 表格数={len(doc.tables)}"
                 f" | 输出长度: {len(full_text)} 字符")

    return full_text


def analyze_docx_from_base64(base64_data: str, filename: str = "document.docx", user_text: str = "") -> Optional[str]:
    """
    从 base64 编码数据解析 .docx 文件。

    Args:
        base64_data: base64 编码的文件内容
        filename: 原始文件名（用于推断扩展名）
        user_text: 用户针对此文档的问题

    Returns:
        提取的文档文本内容；失败返回 None
    """
    import base64
    import tempfile

    try:
        raw_bytes = base64.b64decode(base64_data)
    except Exception as e:
        print(f"[word_analyzer] base64 解码失败: {e}")
        return None

    # 写入临时文件再解析（python-docx 需要文件路径）
    tmp_dir = tempfile.gettempdir()
    safe_filename = filename.replace("/", "_").replace("\\", "_") or "document.docx"
    tmp_path = os.path.join(tmp_dir, f"docx_{hash(base64_data[:16])}_{safe_filename}")

    try:
        with open(tmp_path, "wb") as f:
            f.write(raw_bytes)

        result = analyze_docx(tmp_path, user_text)
        return result
    finally:
        # 清理临时文件
        if os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except OSError:
                pass


def _try_antiword(doc_path: str) -> Optional[str]:
    """降级方案：尝试用 antiword 提取 .doc 纯文本。"""
    import subprocess
    import shutil

    antiword = shutil.which("antiword")
    if not antiword:
        print("[word_analyzer] antiword 未安装且 python-docx 不可用")
        return None

    try:
        result = subprocess.run(
            [antiword, doc_path],
            capture_output=True,
            text=True,
            timeout=30,
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
        else:
            print(f"[word_analyzer] antiword 执行失败: {result.stderr}")
            return None
    except Exception as e:
        print(f"[word_analyzer] antiword 异常: {e}")
        return None
